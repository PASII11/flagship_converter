"""Конвертер документов (PDF, DOCX, MD) на базе Docling и pdfkit."""
from __future__ import annotations

import base64
import gc
import html
import mimetypes
import os
import shutil
import subprocess
import sys
import tempfile
import time
from collections.abc import Callable
from pathlib import Path
from typing import TYPE_CHECKING

import markdown
import pdfkit
from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

if TYPE_CHECKING:
    from docling_core.types.doc.document import DoclingDocument

# Порядок ниже намеренный: huggingface_hub читает эти переменные окружения
# при импорте, поэтому env-блок должен стоять до import huggingface_hub
# (см. noqa: E402/I001 ниже). Не переставлять ради линтера.
os.environ["HF_HUB_DISABLE_SYMLINKS_WARNING"] = "1"
os.environ["HUGGINGFACE_HUB_VERBOSITY"] = "error"
os.environ.setdefault("HF_HUB_OFFLINE", "1")
os.environ.setdefault("TRANSFORMERS_OFFLINE", "1")
os.environ["LOKY_MAX_CPU_COUNT"] = "1"
os.environ["OMP_NUM_THREADS"] = "1"

import huggingface_hub.file_download as _hf_fd  # noqa: E402, I001

from flagship_converter.core.converters.base import safe_output_path  # noqa: E402
from flagship_converter.core.converters.media import get_binary_path, get_wkhtmltopdf_path  # noqa: E402
from flagship_converter.core.converters.pdf_docx import (  # noqa: E402
    convert_pdf_to_docx,
    sanitize_for_xml,
)


def _safe_symlink(src: str, dst: str, **kwargs: object) -> None:
    try:
        os.symlink(src, dst)
    except (OSError, NotImplementedError):
        src_abs = os.path.join(os.path.dirname(dst), src) if not os.path.isabs(src) else src
        if not os.path.exists(dst):
            shutil.copy2(src_abs, dst)

_hf_fd._create_symlink = _safe_symlink

SUPPORTED_INPUT = {".pdf", ".docx", ".md"}
SUPPORTED_OUTPUT = {"pdf", "docx", "md"}


def _resource_path(name: str) -> Path:
    if getattr(sys, "frozen", False) and hasattr(sys, "_MEIPASS"):
        return Path(sys._MEIPASS) / name
    return Path.cwd() / "build_tools" / name


def _optional_binary_path(name: str, win_default_paths: list[str] | None = None) -> str | None:
    try:
        return get_binary_path(name, win_default_paths=win_default_paths)
    except RuntimeError:
        return None


def _docling_artifacts_path() -> Path | None:
    candidates = [_resource_path("docling_models")]
    try:
        repo_root = Path(__file__).resolve().parents[4]
        candidates.append(repo_root / "build_tools" / "docling_models")
    except IndexError:
        pass

    for candidate in candidates:
        if candidate.exists() and candidate.is_dir():
            return candidate

    return None


def _docling_convert_document(
    input_path: Path, generate_images: bool = False
) -> DoclingDocument:
    artifacts_path = _docling_artifacts_path()
    if artifacts_path is None:
        raise RuntimeError("Docling offline models are not bundled.")

    from docling.datamodel.accelerator_options import AcceleratorOptions
    from docling.datamodel.base_models import InputFormat
    from docling.datamodel.pipeline_options import PdfPipelineOptions
    from docling.document_converter import DocumentConverter, PdfFormatOption

    pipeline_options = PdfPipelineOptions()
    pipeline_options.accelerator_options = AcceleratorOptions(num_threads=1)
    pipeline_options.artifacts_path = artifacts_path
    pipeline_options.enable_remote_services = False
    if generate_images:
        pipeline_options.generate_picture_images = True
        pipeline_options.images_scale = 2.0

    converter = None
    conv_res = None
    try:
        converter = DocumentConverter(
            allowed_formats=[InputFormat.PDF, InputFormat.DOCX],
            format_options={
                InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options)
            },
        )
        conv_res = converter.convert(str(input_path))
        if conv_res.document is None:
            raise RuntimeError(f"Docling returned no document for {input_path.name}")
        document = conv_res.document
        return document
    finally:
        del conv_res
        del converter
        gc.collect()


def _extract_pdf_text_as_markdown(input_path: Path) -> str:
    import pypdfium2 as pdfium

    pdf = pdfium.PdfDocument(str(input_path))
    pages: list[str] = []
    try:
        for index in range(len(pdf)):
            page = pdf[index]
            try:
                textpage = page.get_textpage()
                try:
                    text = textpage.get_text_range().strip()
                finally:
                    textpage.close()
            finally:
                page.close()

            if text:
                pages.append(f"## Page {index + 1}\n\n{text}")
    finally:
        pdf.close()

    if not pages:
        raise RuntimeError(
            "PDF contains no extractable text. Bundle Docling offline models to process "
            "scanned PDFs or image-only documents."
        )

    return "\n\n".join(pages)


def _extract_docx_as_markdown(input_path: Path) -> str:
    source_doc = Document(str(input_path))
    blocks: list[str] = []

    for paragraph in source_doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)

    for table in source_doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                blocks.append(" | ".join(cells))

    if not blocks:
        raise RuntimeError(f"Document contains no extractable text: {input_path.name}")

    return "\n\n".join(blocks)


def _iter_docx_blocks(document: DocxDocument):
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _image_data_uri(document: DocxDocument, embed_id: str) -> str | None:
    part = document.part.related_parts.get(embed_id)
    if not part:
        return None

    content_type = getattr(part, "content_type", None)
    if not content_type:
        content_type = mimetypes.guess_type(getattr(part, "partname", ""))[0] or "image/png"

    encoded = base64.b64encode(part.blob).decode("ascii")
    return f"data:{content_type};base64,{encoded}"


def _run_to_html(document: DocxDocument, run) -> str:
    pieces: list[str] = []
    for drawing in run._element.xpath(".//a:blip"):
        embed_id = drawing.get(qn("r:embed"))
        if not embed_id:
            continue
        data_uri = _image_data_uri(document, embed_id)
        if data_uri:
            pieces.append(f'<img src="{data_uri}" alt="" />')

    text = html.escape(run.text).replace("\n", "<br>")
    if text:
        styles: list[str] = []
        if run.font.size:
            styles.append(f"font-size:{run.font.size.pt:.1f}pt")
        if run.font.color and run.font.color.rgb:
            styles.append(f"color:#{run.font.color.rgb}")

        open_tags: list[str] = []
        close_tags: list[str] = []
        if run.bold:
            open_tags.append("<strong>")
            close_tags.insert(0, "</strong>")
        if run.italic:
            open_tags.append("<em>")
            close_tags.insert(0, "</em>")
        if run.underline:
            open_tags.append("<u>")
            close_tags.insert(0, "</u>")

        style_attr = f' style="{";".join(styles)}"' if styles else ""
        pieces.append(f"<span{style_attr}>{''.join(open_tags)}{text}{''.join(close_tags)}</span>")

    return "".join(pieces)


def _paragraph_to_html(document: DocxDocument, paragraph: Paragraph) -> str:
    content = "".join(_run_to_html(document, run) for run in paragraph.runs)
    if not content.strip():
        return ""

    styles: list[str] = []
    if paragraph.alignment is not None:
        align_map = {0: "left", 1: "center", 2: "right", 3: "justify"}
        if paragraph.alignment in align_map:
            styles.append(f"text-align:{align_map[paragraph.alignment]}")

    style_attr = f' style="{";".join(styles)}"' if styles else ""
    style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
    if style_name.startswith("heading"):
        level = next((ch for ch in style_name if ch.isdigit()), "1")
        level_int = max(1, min(int(level), 6))
        return f"<h{level_int}{style_attr}>{content}</h{level_int}>"

    return f"<p{style_attr}>{content}</p>"


def _table_to_html(document: DocxDocument, table: Table) -> str:
    rows: list[str] = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cell_html = "".join(_paragraph_to_html(document, p) for p in cell.paragraphs)
            cells.append(f"<td>{cell_html}</td>")
        rows.append(f"<tr>{''.join(cells)}</tr>")
    return f"<table>{''.join(rows)}</table>"


def _docx_to_html(input_path: Path) -> str:
    document = Document(str(input_path))
    body: list[str] = []
    for block in _iter_docx_blocks(document):
        if isinstance(block, Paragraph):
            body.append(_paragraph_to_html(document, block))
        elif isinstance(block, Table):
            body.append(_table_to_html(document, block))

    return f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
@page {{ margin: 18mm; }}
body {{ font-family: "Calibri", "Arial", sans-serif; line-height: 1.35; color: #111; }}
p {{ margin: 0 0 8pt; }}
h1, h2, h3, h4, h5, h6 {{ margin: 12pt 0 6pt; page-break-after: avoid; }}
img {{ max-width: 100%; height: auto; }}
table {{ border-collapse: collapse; width: 100%; margin: 8pt 0; }}
td, th {{ border: 1px solid #999; padding: 4pt; vertical-align: top; }}
</style>
</head>
<body>
{''.join(body)}
</body>
</html>"""


def _find_soffice_path() -> str | None:
    return _optional_binary_path(
        "soffice",
        win_default_paths=[
            r"C:\Program Files\LibreOffice\program",
            r"C:\Program Files (x86)\LibreOffice\program",
        ],
    )


def _run_cancellable_process(cmd: list[str], cancel_cb: Callable[[], bool]) -> None:
    creationflags = getattr(subprocess, "CREATE_NO_WINDOW", 0) if sys.platform == "win32" else 0
    process = subprocess.Popen(
        cmd,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.PIPE,
        text=True,
        encoding="utf-8",
        errors="replace",
        creationflags=creationflags,
    )
    stderr = ""
    while process.poll() is None:
        if cancel_cb():
            process.terminate()
            try:
                process.wait(timeout=5)
            except subprocess.TimeoutExpired:
                process.kill()
                process.wait(timeout=5)
            return
        time.sleep(0.1)

    if process.stderr:
        stderr = process.stderr.read()

    if process.returncode != 0:
        raise RuntimeError(f"LibreOffice conversion failed: {stderr[-1000:]}")


def _convert_docx_to_pdf_with_soffice(
    input_path: Path,
    output_path: Path,
    cancel_cb: Callable[[], bool],
) -> bool:
    soffice = _find_soffice_path()
    if not soffice:
        return False

    with tempfile.TemporaryDirectory(prefix="flagship-docx-pdf-") as tmp:
        tmp_dir = Path(tmp)
        profile_dir = tmp_dir / "profile"
        profile_uri = profile_dir.as_uri()
        _run_cancellable_process(
            [
                soffice,
                "--headless",
                "--nologo",
                "--nofirststartwizard",
                f"-env:UserInstallation={profile_uri}",
                "--convert-to",
                "pdf",
                "--outdir",
                str(tmp_dir),
                str(input_path),
            ],
            cancel_cb,
        )
        generated = tmp_dir / f"{input_path.stem}.pdf"
        if not generated.exists() or generated.stat().st_size == 0:
            raise RuntimeError("LibreOffice did not produce a PDF output.")
        output_path.parent.mkdir(parents=True, exist_ok=True)
        shutil.move(str(generated), str(output_path))
    return True


def _convert_docx_to_pdf_with_html(input_path: Path, output_path: Path) -> None:
    html_content = _docx_to_html(input_path)
    wk_path = get_wkhtmltopdf_path()
    config = pdfkit.configuration(wkhtmltopdf=wk_path)
    options = {
        "quiet": "",
        "encoding": "UTF-8",
        "enable-local-file-access": "",
        "page-size": "A4",
    }
    output_path.parent.mkdir(parents=True, exist_ok=True)
    pdfkit.from_string(html_content, str(output_path), configuration=config, options=options)


def _convert_docx_to_pdf_preserving_layout(
    input_path: Path,
    output_path: Path,
    cancel_cb: Callable[[], bool],
) -> None:
    if _convert_docx_to_pdf_with_soffice(input_path, output_path, cancel_cb):
        return
    if cancel_cb():
        return
    _convert_docx_to_pdf_with_html(input_path, output_path)


class DocConverter:
    """Конвертирует документы с сохранением структуры (PyInstaller-safe)."""

    supported_outputs = SUPPORTED_OUTPUT

    def __init__(self) -> None:
        pass

    def can_handle(self, path: Path) -> bool:
        return path.suffix.lower() in SUPPORTED_INPUT

    def build_output_path(
        self,
        input_path: Path,
        output_dir: Path,
        target_ext: str,
        overwrite: bool,
    ) -> Path:
        output_dir.mkdir(parents=True, exist_ok=True)
        base = output_dir / f"{input_path.stem}.{target_ext.lstrip('.')}"
        return safe_output_path(base, overwrite)

    def convert(
        self,
        input_path: Path,
        output_path: Path,
        params: dict[str, object],
        cancel_cb: Callable[[], bool],
        progress_cb: Callable[[int], None] | None = None,
    ) -> None:
        if cancel_cb():
            return

        target_ext = output_path.suffix.lower().lstrip(".")

        if progress_cb:
            progress_cb(10)

        if input_path.suffix.lower() == ".docx" and target_ext == "pdf":
            _convert_docx_to_pdf_preserving_layout(input_path, output_path, cancel_cb)
            if progress_cb:
                progress_cb(100)
            return

        if input_path.suffix.lower() == ".pdf" and target_ext == "docx":
            docling_factory = None
            if _docling_artifacts_path() is not None:

                def _factory(path: Path) -> DoclingDocument:
                    return _docling_convert_document(path, generate_images=True)

                docling_factory = _factory
            convert_pdf_to_docx(
                input_path,
                output_path,
                cancel_cb=cancel_cb,
                progress_cb=progress_cb,
                docling_factory=docling_factory,
            )
            return

        if input_path.suffix.lower() == ".md":
            md_text = input_path.read_text(encoding="utf-8")
        else:
            artifacts_path = _docling_artifacts_path()
            if artifacts_path is None:
                if input_path.suffix.lower() == ".pdf":
                    md_text = _extract_pdf_text_as_markdown(input_path)
                elif input_path.suffix.lower() == ".docx":
                    md_text = _extract_docx_as_markdown(input_path)
                else:
                    raise RuntimeError(f"Unsupported document source format: {input_path.suffix}")
                if progress_cb:
                    progress_cb(50)
                if cancel_cb():
                    return
            else:
                md_text = _docling_convert_document(input_path).export_to_markdown()

        if cancel_cb():
            return

        if progress_cb:
            progress_cb(60)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        if target_ext == "md":
            output_path.write_text(md_text, encoding="utf-8")

        elif target_ext == "pdf":
            html_content = markdown.markdown(md_text, extensions=["tables"])
            full_html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><style>
body {{ font-family: sans-serif; line-height: 1.5; padding: 20px; }}
table {{ border-collapse: collapse; width: 100%; margin-bottom: 20px; }}
th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
th {{ background-color: #f2f2f2; }}
img {{ max-width: 100%; }}
</style></head><body>{html_content}</body></html>"""

            wk_path = get_wkhtmltopdf_path()
            config = pdfkit.configuration(wkhtmltopdf=wk_path)
            options = {"quiet": "", "encoding": "UTF-8", "enable-local-file-access": ""}
            pdfkit.from_string(full_html, str(output_path), configuration=config, options=options)

        elif target_ext == "docx":
            doc = Document()
            for line in sanitize_for_xml(md_text).split("\n"):
                if line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("### "):
                    doc.add_heading(line[4:], level=3)
                elif line.strip():
                    doc.add_paragraph(line)
            doc.save(str(output_path))

        else:
            raise RuntimeError(f"Unsupported document target format: {target_ext}")

        if progress_cb:
            progress_cb(100)
