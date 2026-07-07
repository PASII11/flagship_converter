"""PDF → DOCX конвертация с сохранением оформления (pdf2docx + Docling)."""

from __future__ import annotations

import io
import logging
import re
from collections.abc import Callable
from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document
from docx.shared import Inches

if TYPE_CHECKING:
    from docling_core.types.doc.document import DoclingDocument

logger = logging.getLogger(__name__)

_XML_INCOMPATIBLE = re.compile(
    "[\\x00-\\x08\\x0b\\x0c\\x0e-\\x1f\\x7f"
    "\\ud800-\\udfff\\ufdd0-\\ufddf\\ufffe\\uffff]"
)


def sanitize_for_xml(text: str) -> str:
    return _XML_INCOMPATIBLE.sub("", text)


def pdf_has_text_layer(
    input_path: Path,
    min_chars_per_page: int = 25,
    min_text_page_ratio: float = 0.5,
) -> bool:
    """Цифровой PDF: не меньше половины страниц дают > 25 извлекаемых символов."""
    import pypdfium2 as pdfium

    pdf = pdfium.PdfDocument(str(input_path))
    try:
        page_count = len(pdf)
        if page_count == 0:
            return False
        text_pages = 0
        for index in range(page_count):
            page = pdf[index]
            try:
                textpage = page.get_textpage()
                try:
                    text = textpage.get_text_range().strip()
                finally:
                    textpage.close()
            finally:
                page.close()
            if len(text) > min_chars_per_page:
                text_pages += 1
        return text_pages / page_count >= min_text_page_ratio
    finally:
        pdf.close()


def convert_with_pdf2docx(input_path: Path, output_path: Path) -> None:
    """Цифровой путь: pdf2docx воспроизводит раскладку, шрифты, картинки, таблицы."""
    from pdf2docx import Converter

    logging.getLogger("pdf2docx").setLevel(logging.ERROR)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    converter = Converter(str(input_path))
    try:
        converter.convert(str(output_path))
    finally:
        converter.close()

    if not output_path.exists() or output_path.stat().st_size == 0:
        raise RuntimeError("pdf2docx did not produce a DOCX output.")


def docling_document_to_docx(dl_doc: DoclingDocument, output_path: Path) -> None:
    """Скан-путь: строит DOCX напрямую из структуры DoclingDocument."""
    from docling_core.types.doc import DocItemLabel
    from docling_core.types.doc.document import (
        ListItem,
        PictureItem,
        SectionHeaderItem,
        TableItem,
        TextItem,
    )

    doc = Document()
    blocks_added = 0

    for item, _level in dl_doc.iterate_items():
        if isinstance(item, TableItem):
            grid = item.data.grid
            if not grid:
                continue
            table = doc.add_table(rows=len(grid), cols=max(len(row) for row in grid))
            table.style = "Table Grid"
            for row_index, row in enumerate(grid):
                for col_index, cell in enumerate(row):
                    docx_cell = table.rows[row_index].cells[col_index]
                    docx_cell.text = sanitize_for_xml(cell.text)
            blocks_added += 1
        elif isinstance(item, PictureItem):
            image = item.get_image(dl_doc)
            if image is None:
                continue
            buffer = io.BytesIO()
            image.save(buffer, format="PNG")
            buffer.seek(0)
            doc.add_picture(buffer, width=Inches(5.5))
            blocks_added += 1
        elif isinstance(item, SectionHeaderItem):
            doc.add_heading(sanitize_for_xml(item.text), level=min(item.level, 6))
            blocks_added += 1
        elif isinstance(item, ListItem):
            style = "List Number" if item.enumerated else "List Bullet"
            doc.add_paragraph(sanitize_for_xml(item.text), style=style)
            blocks_added += 1
        elif isinstance(item, TextItem):
            if item.label in (DocItemLabel.PAGE_HEADER, DocItemLabel.PAGE_FOOTER):
                continue
            text = sanitize_for_xml(item.text)
            if not text.strip():
                continue
            if item.label == DocItemLabel.TITLE:
                doc.add_heading(text, level=0)
            elif item.label == DocItemLabel.CAPTION:
                paragraph = doc.add_paragraph()
                paragraph.add_run(text).italic = True
            else:
                doc.add_paragraph(text)
            blocks_added += 1

    if blocks_added == 0:
        raise RuntimeError("Docling document contains no renderable content.")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def extract_text_paragraphs(input_path: Path) -> list[str]:
    """Аварийный путь: сырой текст PDF со склейкой перенесённых строк в абзацы."""
    import pypdfium2 as pdfium

    lines: list[str] = []
    pdf = pdfium.PdfDocument(str(input_path))
    try:
        for index in range(len(pdf)):
            page = pdf[index]
            try:
                textpage = page.get_textpage()
                try:
                    text = textpage.get_text_range()
                finally:
                    textpage.close()
            finally:
                page.close()
            lines.extend(text.splitlines())
            lines.append("")
    finally:
        pdf.close()

    paragraphs: list[str] = []
    current: list[str] = []
    for line in lines:
        stripped = line.strip()
        if stripped:
            current.append(stripped)
        elif current:
            paragraphs.append(" ".join(current))
            current = []
    if current:
        paragraphs.append(" ".join(current))

    if not paragraphs:
        raise RuntimeError(
            "PDF contains no extractable text. Bundle Docling offline models to process "
            "scanned PDFs or image-only documents."
        )
    return paragraphs


DoclingFactory = Callable[[Path], "DoclingDocument"]


def convert_pdf_to_docx(
    input_path: Path,
    output_path: Path,
    cancel_cb: Callable[[], bool],
    progress_cb: Callable[[int], None] | None = None,
    docling_factory: DoclingFactory | None = None,
) -> None:
    """Цепочка: pdf2docx (цифровые) → Docling (сканы) → текст с абзацами."""
    if cancel_cb():
        return

    try:
        is_digital = pdf_has_text_layer(input_path)
    except Exception as exc:
        logger.warning("PDF text-layer detection failed for %s: %s", input_path.name, exc)
        is_digital = True

    if progress_cb:
        progress_cb(10)
    if cancel_cb():
        return

    last_error: Exception | None = None

    if is_digital:
        try:
            if progress_cb:
                progress_cb(20)
            convert_with_pdf2docx(input_path, output_path)
            if progress_cb:
                progress_cb(100)
            return
        except Exception as exc:
            last_error = exc
            logger.warning("pdf2docx failed for %s, falling back: %s", input_path.name, exc)

    if cancel_cb():
        return

    if docling_factory is not None:
        try:
            if progress_cb:
                progress_cb(40)
            dl_doc = docling_factory(input_path)
            if cancel_cb():
                return
            if progress_cb:
                progress_cb(90)
            docling_document_to_docx(dl_doc, output_path)
            if progress_cb:
                progress_cb(100)
            return
        except Exception as exc:
            last_error = exc
            logger.warning("Docling failed for %s, falling back: %s", input_path.name, exc)

    if cancel_cb():
        return

    try:
        paragraphs = extract_text_paragraphs(input_path)
    except Exception as exc:
        raise RuntimeError(
            f"Failed to convert {input_path.name} to DOCX: {last_error or exc}"
        ) from exc

    doc = Document()
    for paragraph in paragraphs:
        doc.add_paragraph(sanitize_for_xml(paragraph))
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    if progress_cb:
        progress_cb(100)
