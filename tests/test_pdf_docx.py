"""Тесты PDF → DOCX конвертации с сохранением оформления."""

from __future__ import annotations

import io
import zipfile
from pathlib import Path

import pytest
from PIL import Image


def _make_digital_pdf(
    path: Path,
    paragraphs: list[str] | None = None,
    with_image: bool = False,
) -> None:
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    y = 72.0
    lines = paragraphs or ["Hello PDF world, this is a digital document for testing."]
    for text in lines:
        page.insert_text((72, y), text, fontsize=12)
        y += 20
    if with_image:
        image = Image.new("RGB", (120, 60), (30, 120, 220))
        buffer = io.BytesIO()
        image.save(buffer, format="PNG")
        page.insert_image(fitz.Rect(72, y, 192, y + 60), stream=buffer.getvalue())
    doc.save(str(path))
    doc.close()


def _make_image_only_pdf(path: Path) -> None:
    import fitz

    image = Image.new("RGB", (200, 100), (200, 30, 30))
    buffer = io.BytesIO()
    image.save(buffer, format="PNG")
    doc = fitz.open()
    page = doc.new_page(width=300, height=200)
    page.insert_image(fitz.Rect(50, 50, 250, 150), stream=buffer.getvalue())
    doc.save(str(path))
    doc.close()


def test_sanitize_for_xml_strips_control_chars() -> None:
    from flagship_converter.core.converters.pdf_docx import sanitize_for_xml

    assert sanitize_for_xml("ok\x00\x0btext") == "oktext"


def test_pdf_has_text_layer_true_for_digital(tmp_path: Path) -> None:
    from flagship_converter.core.converters.pdf_docx import pdf_has_text_layer

    pdf = tmp_path / "digital.pdf"
    _make_digital_pdf(pdf)

    assert pdf_has_text_layer(pdf) is True


def test_pdf_has_text_layer_false_for_image_only(tmp_path: Path) -> None:
    from flagship_converter.core.converters.pdf_docx import pdf_has_text_layer

    pdf = tmp_path / "scan.pdf"
    _make_image_only_pdf(pdf)

    assert pdf_has_text_layer(pdf) is False


def test_convert_with_pdf2docx_preserves_text_and_images(tmp_path: Path) -> None:
    from docx import Document

    from flagship_converter.core.converters.pdf_docx import convert_with_pdf2docx

    pdf = tmp_path / "digital.pdf"
    out = tmp_path / "digital.docx"
    _make_digital_pdf(pdf, with_image=True)

    convert_with_pdf2docx(pdf, out)

    assert out.exists()
    assert out.stat().st_size > 0
    text = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "digital document" in text
    media = [n for n in zipfile.ZipFile(out).namelist() if n.startswith("word/media/")]
    assert media, "embedded PDF image must survive conversion"


def test_docling_document_to_docx_renders_structure(tmp_path: Path) -> None:
    from docling_core.types.doc import DocItemLabel, DoclingDocument
    from docling_core.types.doc.document import ImageRef, TableCell, TableData
    from docx import Document

    from flagship_converter.core.converters.pdf_docx import docling_document_to_docx

    dl_doc = DoclingDocument(name="sample")
    dl_doc.add_heading("Section title", level=1)
    dl_doc.add_text(label=DocItemLabel.PARAGRAPH, text="First paragraph")
    dl_doc.add_list_item("Bullet one")
    cells = [
        TableCell(
            text="A1",
            start_row_offset_idx=0,
            end_row_offset_idx=1,
            start_col_offset_idx=0,
            end_col_offset_idx=1,
        ),
        TableCell(
            text="B1",
            start_row_offset_idx=0,
            end_row_offset_idx=1,
            start_col_offset_idx=1,
            end_col_offset_idx=2,
        ),
    ]
    dl_doc.add_table(data=TableData(table_cells=cells, num_rows=1, num_cols=2))
    dl_doc.add_picture(
        image=ImageRef.from_pil(Image.new("RGB", (16, 16), (255, 0, 0)), dpi=72)
    )

    out = tmp_path / "rendered.docx"
    docling_document_to_docx(dl_doc, out)

    rendered = Document(str(out))
    styles = [p.style.name for p in rendered.paragraphs]
    texts = [p.text for p in rendered.paragraphs]
    assert any(name.startswith("Heading") for name in styles)
    assert "First paragraph" in texts
    assert "List Bullet" in styles
    assert rendered.tables
    assert rendered.tables[0].rows[0].cells[0].text == "A1"
    assert rendered.tables[0].rows[0].cells[1].text == "B1"
    media = [n for n in zipfile.ZipFile(out).namelist() if n.startswith("word/media/")]
    assert media, "picture must be embedded in the DOCX"


def test_docling_document_to_docx_rejects_empty_document(tmp_path: Path) -> None:
    from docling_core.types.doc import DoclingDocument

    from flagship_converter.core.converters.pdf_docx import docling_document_to_docx

    with pytest.raises(RuntimeError, match="no renderable content"):
        docling_document_to_docx(DoclingDocument(name="empty"), tmp_path / "empty.docx")


def test_extract_text_paragraphs_merges_wrapped_lines(tmp_path: Path) -> None:
    from flagship_converter.core.converters.pdf_docx import extract_text_paragraphs

    pdf = tmp_path / "wrapped.pdf"
    _make_digital_pdf(
        pdf,
        paragraphs=[
            "First line of a paragraph that",
            "continues on the second line",
            "and ends on the third one.",
        ],
    )

    paragraphs = extract_text_paragraphs(pdf)

    assert len(paragraphs) == 1
    assert "that continues on the second line and ends" in paragraphs[0]


def test_extract_text_paragraphs_raises_for_image_only(tmp_path: Path) -> None:
    from flagship_converter.core.converters.pdf_docx import extract_text_paragraphs

    pdf = tmp_path / "scan.pdf"
    _make_image_only_pdf(pdf)

    with pytest.raises(RuntimeError, match="no extractable text"):
        extract_text_paragraphs(pdf)


def test_convert_pdf_to_docx_prefers_pdf2docx_for_digital(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    from flagship_converter.core.converters import pdf_docx

    calls: list[str] = []
    progress: list[int] = []

    def fake_pdf2docx(inp: Path, out: Path) -> None:
        calls.append("pdf2docx")
        out.write_bytes(b"stub-docx")

    monkeypatch.setattr(pdf_docx, "pdf_has_text_layer", lambda _p: True)
    monkeypatch.setattr(pdf_docx, "convert_with_pdf2docx", fake_pdf2docx)

    pdf_docx.convert_pdf_to_docx(
        tmp_path / "in.pdf",
        tmp_path / "out.docx",
        cancel_cb=lambda: False,
        progress_cb=progress.append,
    )

    assert calls == ["pdf2docx"]
    assert progress[-1] == 100


def test_convert_pdf_to_docx_falls_back_to_docling(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    from flagship_converter.core.converters import pdf_docx

    calls: list[str] = []
    sentinel = object()

    def failing_pdf2docx(inp: Path, out: Path) -> None:
        calls.append("pdf2docx")
        raise RuntimeError("boom")

    def fake_renderer(dl_doc: object, out: Path) -> None:
        assert dl_doc is sentinel
        calls.append("docling")
        out.write_bytes(b"stub-docx")

    monkeypatch.setattr(pdf_docx, "pdf_has_text_layer", lambda _p: True)
    monkeypatch.setattr(pdf_docx, "convert_with_pdf2docx", failing_pdf2docx)
    monkeypatch.setattr(pdf_docx, "docling_document_to_docx", fake_renderer)

    pdf_docx.convert_pdf_to_docx(
        tmp_path / "in.pdf",
        tmp_path / "out.docx",
        cancel_cb=lambda: False,
        docling_factory=lambda _p: sentinel,  # type: ignore[arg-type,return-value]
    )

    assert calls == ["pdf2docx", "docling"]


def test_convert_pdf_to_docx_last_resort_writes_paragraphs(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    from docx import Document

    from flagship_converter.core.converters import pdf_docx

    monkeypatch.setattr(pdf_docx, "pdf_has_text_layer", lambda _p: False)
    monkeypatch.setattr(
        pdf_docx, "extract_text_paragraphs", lambda _p: ["First paragraph", "Second one"]
    )

    out = tmp_path / "out.docx"
    pdf_docx.convert_pdf_to_docx(
        tmp_path / "in.pdf", out, cancel_cb=lambda: False, docling_factory=None
    )

    texts = [p.text for p in Document(str(out)).paragraphs]
    assert "First paragraph" in texts
    assert "Second one" in texts


def test_convert_pdf_to_docx_cancelled_before_start(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    from flagship_converter.core.converters import pdf_docx

    def must_not_run(_p: Path) -> bool:
        raise AssertionError("classifier must not run after cancellation")

    monkeypatch.setattr(pdf_docx, "pdf_has_text_layer", must_not_run)

    out = tmp_path / "out.docx"
    pdf_docx.convert_pdf_to_docx(tmp_path / "in.pdf", out, cancel_cb=lambda: True)

    assert not out.exists()


def test_convert_pdf_to_docx_raises_when_all_engines_fail(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    from flagship_converter.core.converters import pdf_docx

    def failing_extract(_p: Path) -> list[str]:
        raise RuntimeError("no extractable text")

    monkeypatch.setattr(pdf_docx, "pdf_has_text_layer", lambda _p: False)
    monkeypatch.setattr(pdf_docx, "extract_text_paragraphs", failing_extract)

    with pytest.raises(RuntimeError, match="Failed to convert in.pdf"):
        pdf_docx.convert_pdf_to_docx(
            tmp_path / "in.pdf",
            tmp_path / "out.docx",
            cancel_cb=lambda: False,
            docling_factory=None,
        )
