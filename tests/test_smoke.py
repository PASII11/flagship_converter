"""Smoke-тесты: проверка импортов и базовой логики без GUI."""

from __future__ import annotations

import io
import os
import tempfile
from pathlib import Path

import pytest
from PIL import Image


def test_import_package() -> None:
    import flagship_converter
    assert flagship_converter.__version__ == "0.1.0"


def test_import_app_main() -> None:
    from flagship_converter.app import main
    assert callable(main)


def test_image_converter_roundtrip() -> None:
    """Создать PNG → конвертировать в JPEG → файл существует и не пустой."""
    from flagship_converter.core.converters.image import ImageConverter

    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / "test.png"
        out = Path(tmp) / "out" / "test.jpg"

        img = Image.new("RGB", (64, 64), color=(120, 80, 200))
        img.save(src)

        converter = ImageConverter()
        assert converter.can_handle(src)
        out.parent.mkdir(parents=True, exist_ok=True)
        converter.convert(src, out, {"quality": 85}, cancel_cb=lambda: False)

        assert out.exists()
        assert out.stat().st_size > 0


def test_safe_output_path_no_collision() -> None:
    from flagship_converter.core.converters.base import safe_output_path

    with tempfile.TemporaryDirectory() as tmp:
        p = Path(tmp) / "file.jpg"
        result = safe_output_path(p, overwrite=False)
        assert result == p  # файла нет — возвращаем как есть


def test_safe_output_path_collision() -> None:
    from flagship_converter.core.converters.base import safe_output_path

    with tempfile.TemporaryDirectory() as tmp:
        p = Path(tmp) / "file.jpg"
        p.touch()  # имитируем существующий файл
        result = safe_output_path(p, overwrite=False)
        assert result.name == "file_1.jpg"


def test_build_job_rejects_unsupported_target() -> None:
    from flagship_converter.core.engine import ConversionEngine

    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / "image.png"
        src.touch()

        engine = ConversionEngine()
        job = engine.build_job(src, Path(tmp) / "out", "mp3", False, {})

        assert job is None


class FakeConverter:
    supported_outputs = {"out"}

    def __init__(self, mode: str = "success") -> None:
        self.mode = mode

    def can_handle(self, path: Path) -> bool:
        return path.suffix == ".in"

    def build_output_path(
        self,
        input_path: Path,
        output_dir: Path,
        target_ext: str,
        overwrite: bool,
    ) -> Path:
        from flagship_converter.core.converters.base import safe_output_path

        output_dir.mkdir(parents=True, exist_ok=True)
        return safe_output_path(output_dir / f"{input_path.stem}.{target_ext}", overwrite)

    def convert(self, input_path, output_path, params, cancel_cb, progress_cb=None) -> None:
        from flagship_converter.core.converters.base import ConversionCancelled

        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text("converted", encoding="utf-8")
        if self.mode == "cancel":
            raise ConversionCancelled()
        if self.mode == "fail":
            raise RuntimeError("fake failure")
        if progress_cb:
            progress_cb(40)


def test_build_plan_reserves_duplicate_output_names() -> None:
    from flagship_converter.core.engine import ConversionEngine

    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        src_a = root / "a" / "same.in"
        src_b = root / "b" / "same.in"
        src_a.parent.mkdir()
        src_b.parent.mkdir()
        src_a.touch()
        src_b.touch()

        engine = ConversionEngine()
        engine._converters = [FakeConverter()]  # type: ignore[attr-defined]

        plan = engine.build_plan(
            [src_a, src_b],
            root / "out",
            "out",
            overwrite=False,
            quality=85,
            lossless_webp=False,
            audio_bitrate="192k",
            video_bitrate="2.5M",
            video_codec="CPU",
        )

        assert len(plan.jobs) == 2
        assert plan.jobs[0].output_path != plan.jobs[1].output_path


def test_execute_plan_publishes_temp_output_atomically() -> None:
    from flagship_converter.core.engine import ConversionEngine
    from flagship_converter.core.models import ConversionJob, ConversionPlan, JobStatus

    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        src = root / "input.in"
        final = root / "out" / "input.out"
        src.touch()
        job = ConversionJob(src, final, "FakeConverter", {}, "out", overwrite=False)
        plan = ConversionPlan([job])
        finished: list[str] = []
        progress: list[int] = []

        engine = ConversionEngine()
        engine._converters = [FakeConverter()]  # type: ignore[attr-defined]
        engine.execute_plan(
            plan,
            cancel_cb=lambda: False,
            on_job_started=lambda _job_id: None,
            on_job_finished=finished.append,
            on_job_failed=lambda _job_id, error: pytest.fail(error),
            on_job_progress=lambda _job_id, percent: progress.append(percent),
        )

        assert job.status == JobStatus.DONE
        assert final.read_text(encoding="utf-8") == "converted"
        assert not list(final.parent.glob("*.part.*"))
        assert finished == [job.id]
        assert progress[-1] == 100


def test_execute_plan_cleans_temp_output_on_cancel() -> None:
    from flagship_converter.core.engine import ConversionEngine
    from flagship_converter.core.models import ConversionJob, ConversionPlan, JobStatus

    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        src = root / "input.in"
        final = root / "out" / "input.out"
        src.touch()
        job = ConversionJob(src, final, "FakeConverter", {}, "out", overwrite=False)
        cancelled: list[str] = []

        engine = ConversionEngine()
        engine._converters = [FakeConverter("cancel")]  # type: ignore[attr-defined]
        engine.execute_plan(
            ConversionPlan([job]),
            cancel_cb=lambda: False,
            on_job_started=lambda _job_id: None,
            on_job_finished=lambda _job_id: pytest.fail("cancelled job finished"),
            on_job_failed=lambda _job_id, error: pytest.fail(error),
            on_job_cancelled=cancelled.append,
        )

        assert job.status == JobStatus.CANCELLED
        assert not final.exists()
        assert not list(final.parent.glob("*.part.*"))
        assert cancelled == [job.id]


def test_run_ffmpeg_terminates_process_on_cancel(monkeypatch) -> None:
    from flagship_converter.core.converters import media
    from flagship_converter.core.converters.base import ConversionCancelled

    class FakeProcess:
        def __init__(self) -> None:
            self.stderr = io.StringIO("")
            self.returncode = None
            self.terminated = False
            self.killed = False

        def poll(self):
            return None

        def terminate(self) -> None:
            self.terminated = True

        def kill(self) -> None:
            self.killed = True

        def wait(self, timeout=None):
            self.returncode = -15
            return self.returncode

    process = FakeProcess()
    monkeypatch.setattr(media.subprocess, "Popen", lambda *args, **kwargs: process)

    with pytest.raises(ConversionCancelled):
        media.run_ffmpeg(["ffmpeg", "-version"], cancel_cb=lambda: True)

    assert process.terminated
    assert not process.killed


def test_missing_binary_raises_clear_error(monkeypatch) -> None:
    from flagship_converter.core.converters import media

    monkeypatch.setattr(media.shutil, "which", lambda _name: None)

    with pytest.raises(RuntimeError, match="Required binary"):
        media.get_binary_path("missing-binary-for-test")


def test_docling_artifacts_are_optional_for_text_fallback(monkeypatch, tmp_path) -> None:
    from flagship_converter.core.converters import document

    monkeypatch.setattr(document, "_resource_path", lambda name: tmp_path / name)
    monkeypatch.setattr(
        document,
        "__file__",
        str(
            tmp_path
            / "repo"
            / "src"
            / "flagship_converter"
            / "core"
            / "converters"
            / "document.py"
        ),
    )

    assert os.environ["HF_HUB_OFFLINE"] == "1"
    assert os.environ["TRANSFORMERS_OFFLINE"] == "1"
    assert document._docling_artifacts_path() is None


def test_pdf_to_docx_works_without_docling_models(monkeypatch, tmp_path) -> None:
    from docx import Document

    from flagship_converter.core.converters import document
    from flagship_converter.core.converters.document import DocConverter

    src = tmp_path / "sample.pdf"
    out = tmp_path / "sample.docx"
    src.write_bytes(b"%PDF-1.4\n% test placeholder")

    monkeypatch.setattr(document, "_docling_artifacts_path", lambda: None)
    monkeypatch.setattr(
        document,
        "_extract_pdf_text_as_markdown",
        lambda _path: "# Sample\n\nHello from PDF",
    )

    DocConverter().convert(src, out, {}, cancel_cb=lambda: False)

    assert out.exists()
    assert out.stat().st_size > 0
    text = "\n".join(paragraph.text for paragraph in Document(str(out)).paragraphs)
    assert "Sample" in text
    assert "Hello from PDF" in text


def test_docx_to_html_embeds_images(tmp_path) -> None:
    from docx import Document
    from docx.shared import Inches

    from flagship_converter.core.converters.document import _docx_to_html

    image = tmp_path / "image.png"
    Image.new("RGB", (20, 20), color=(20, 120, 220)).save(image)

    docx_path = tmp_path / "with-image.docx"
    doc = Document()
    doc.add_heading("Image document", level=1)
    doc.add_paragraph("Before image")
    doc.add_picture(str(image), width=Inches(1))
    doc.save(docx_path)

    html_output = _docx_to_html(docx_path)

    assert "Image document" in html_output
    assert "Before image" in html_output
    assert "data:image/png;base64," in html_output


def test_docx_to_pdf_uses_html_fallback_with_images(monkeypatch, tmp_path) -> None:
    from docx import Document
    from docx.shared import Inches

    from flagship_converter.core.converters import document
    from flagship_converter.core.converters.document import DocConverter

    image = tmp_path / "image.png"
    Image.new("RGB", (20, 20), color=(200, 80, 30)).save(image)

    docx_path = tmp_path / "with-image.docx"
    pdf_path = tmp_path / "with-image.pdf"
    doc = Document()
    doc.add_paragraph("Document with embedded image")
    doc.add_picture(str(image), width=Inches(1))
    doc.save(docx_path)

    captured_html: dict[str, str] = {}

    monkeypatch.setattr(document, "_find_soffice_path", lambda: None)
    monkeypatch.setattr(document, "get_wkhtmltopdf_path", lambda: "wkhtmltopdf")
    monkeypatch.setattr(
        document.pdfkit,
        "configuration",
        lambda wkhtmltopdf: {"wkhtmltopdf": wkhtmltopdf},
    )

    def fake_from_string(html_content, output_path, configuration, options):
        captured_html["html"] = html_content
        Path(output_path).write_bytes(b"%PDF-1.4 fake")

    monkeypatch.setattr(document.pdfkit, "from_string", fake_from_string)

    DocConverter().convert(docx_path, pdf_path, {}, cancel_cb=lambda: False)

    assert pdf_path.exists()
    assert "Document with embedded image" in captured_html["html"]
    assert "data:image/png;base64," in captured_html["html"]


def test_video_converter_maps_amd_codec_id(monkeypatch) -> None:
    from flagship_converter.core.converters import video

    captured: dict[str, list[str]] = {}

    def fake_run_ffmpeg(cmd, cancel_cb, progress_cb=None):
        captured["cmd"] = cmd

    monkeypatch.setattr(video, "run_ffmpeg", fake_run_ffmpeg)
    monkeypatch.setattr(video, "get_ffmpeg_path", lambda: "ffmpeg")

    video.VideoConverter().convert(
        Path("in.mp4"), Path("out.mp4"),
        {"video_bitrate": "2.5M", "video_codec": "amd"},
        cancel_cb=lambda: False,
    )
    assert "h264_amf" in captured["cmd"]


def test_video_converter_defaults_to_libx264_for_auto_codec(monkeypatch) -> None:
    from flagship_converter.core.converters import video

    captured: dict[str, list[str]] = {}

    def fake_run_ffmpeg(cmd, cancel_cb, progress_cb=None):
        captured["cmd"] = cmd

    monkeypatch.setattr(video, "run_ffmpeg", fake_run_ffmpeg)
    monkeypatch.setattr(video, "get_ffmpeg_path", lambda: "ffmpeg")

    video.VideoConverter().convert(
        Path("in.mp4"), Path("out.mp4"),
        {"video_bitrate": "2.5M", "video_codec": "auto"},
        cancel_cb=lambda: False,
    )
    assert "libx264" in captured["cmd"]


def test_engine_reports_translated_error_for_missing_converter() -> None:
    from flagship_converter import i18n
    from flagship_converter.core.engine import ConversionEngine
    from flagship_converter.core.models import ConversionJob, ConversionPlan

    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        src = root / "input.in"
        src.touch()
        job = ConversionJob(src, root / "out.out", "GhostConverter", {}, "out", overwrite=False)

        engine = ConversionEngine()
        engine._converters = []  # type: ignore[attr-defined]

        failed: list[tuple[str, str]] = []
        i18n.set_language("en")
        engine.execute_plan(
            ConversionPlan([job]),
            cancel_cb=lambda: False,
            on_job_started=lambda _job_id: None,
            on_job_finished=lambda _job_id: pytest.fail("should not finish"),
            on_job_failed=lambda job_id, error: failed.append((job_id, error)),
        )

        assert failed == [(job.id, "Converter 'GhostConverter' not found")]


def test_image_converter_reports_translated_error_for_corrupt_file(tmp_path: Path) -> None:
    from flagship_converter import i18n
    from flagship_converter.core.converters.image import ImageConverter

    src = tmp_path / "broken.png"
    src.write_bytes(b"not an image")
    out = tmp_path / "out.jpg"

    i18n.set_language("en")
    with pytest.raises(RuntimeError, match="Failed to open image: broken.png"):
        ImageConverter().convert(src, out, {"quality": 85}, cancel_cb=lambda: False)


def test_run_ffmpeg_reports_translated_error_when_stderr_missing(monkeypatch) -> None:
    from flagship_converter import i18n
    from flagship_converter.core.converters import media

    class FakeProcess:
        def __init__(self) -> None:
            self.stderr = None
            self.returncode = None

        def terminate(self) -> None:
            pass

        def wait(self, timeout=None):
            self.returncode = 0
            return self.returncode

    monkeypatch.setattr(media.subprocess, "Popen", lambda *args, **kwargs: FakeProcess())

    i18n.set_language("en")
    with pytest.raises(RuntimeError, match="Failed to open FFmpeg process stderr"):
        media.run_ffmpeg(["ffmpeg", "-version"], cancel_cb=lambda: False)
